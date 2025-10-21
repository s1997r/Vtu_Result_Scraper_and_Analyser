'''

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)


# Load Excel file
file_path = "results.xlsx"  # Replace with actual path
df = pd.read_excel(file_path, sheet_name='Sheet1')

# Clean data
df_clean = df.dropna(how='all')

# Identify relevant columns
subject_name_columns = [col for col in df_clean.columns if col.endswith('_SubjectName')]
result_columns = [col for col in df_clean.columns if col.endswith('_Result')]
total_columns = [col for col in df_clean.columns if col.endswith('_Total')]

# Prepare Subject Code-Name-Appearance Table
subject_table_data = []
serial_no = 1
for sub_col in subject_name_columns:
    subject_code = sub_col.replace('_SubjectName', '')
    subject_name_series = df_clean[sub_col].dropna().astype(str).str.strip()
    subject_name = subject_name_series[subject_name_series != ''].unique()
    if len(subject_name) == 0:
        continue  # skip if blank or NA
    subject_name = subject_name[0]

    # Try to find corresponding _Result column for appearance count
    result_col = f"{subject_code}_Result"
    appeared = df_clean[result_col].notna().sum() if result_col in df_clean.columns else 0

    subject_table_data.append([serial_no, subject_name, subject_code, appeared])
    serial_no += 1

# Calculate Total Full Marks
df_clean['Total_Full_Marks'] = df_clean[total_columns].apply(
    pd.to_numeric, errors='coerce').sum(axis=1)

# Prepare summary results
summary_results = {}
for col in result_columns:
    counts = df_clean[col].value_counts(dropna=True)
    appeared_count = df_clean[col].notna().sum()
    summary_results[col] = {'counts': counts, 'appeared': appeared_count}

# Top 10 students
df_top = df_clean.dropna(subset=['Total_Full_Marks'])
top_10 = df_top.nlargest(10, 'Total_Full_Marks')[[
    'University Seat Number', 'Student Name', 'Total_Full_Marks'
]]

# Create Word document
doc = Document()
doc.add_heading("Student Results Summary Report", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

# Section 0: Subject Code and Name Summary
doc.add_heading("Subject Code and Name Summary", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
table = doc.add_table(rows=1, cols=4)
set_table_borders(table)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "S.No"
hdr_cells[1].text = "Subject Name"
hdr_cells[2].text = "Subject Code"
hdr_cells[3].text = "Total Appeared"

for row in subject_table_data:
    row_cells = table.add_row().cells
    for i, val in enumerate(row):
        row_cells[i].text = str(val)

# Section 1: Summary of Results by Subject
doc.add_heading("Summary of Results by Subject", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
for subject, data in summary_results.items():
    doc.add_heading(subject, level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Result'
    hdr_cells[1].text = 'Count'
    hdr_cells[2].text = 'Percentage'

    total_appeared = data['appeared']
    for result, count in data['counts'].items():
        percentage = (count / total_appeared) * 100 if total_appeared > 0 else 0
        row_cells = table.add_row().cells
        row_cells[0].text = str(result)
        row_cells[1].text = str(count)
        row_cells[2].text = f"{percentage:.2f}%"

    # Total Appeared
    row_cells = table.add_row().cells
    row_cells[0].text = 'Total Appeared'
    row_cells[1].text = str(total_appeared)
    row_cells[2].text = '—'

# Section 2: Top 10 Students
doc.add_heading("Top 10 Students by Total Full Marks", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
table = doc.add_table(rows=1, cols=3)
set_table_borders(table)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'University Seat Number'
hdr_cells[1].text = 'Student Name'
hdr_cells[2].text = 'Total Full Marks'

for _, row in top_10.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(row['University Seat Number'])
    row_cells[1].text = str(row['Student Name'])
    row_cells[2].text = str(int(row['Total_Full_Marks']))

# Section 3: Top Performer in Each Subject
doc.add_heading("Top Performer in Each Subject", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
table = doc.add_table(rows=1, cols=5)
set_table_borders(table)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Student Name"
hdr_cells[1].text = "University Seat Number"
hdr_cells[2].text = "Subject Code"
hdr_cells[3].text = "Subject Marks"
hdr_cells[4].text = "Total Full Marks"

for col in total_columns:
    subject_scores = pd.to_numeric(df_clean[col], errors='coerce')
    max_score = subject_scores.max()
    top_rows = df_clean[subject_scores == max_score]

    for _, row in top_rows.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Student Name'])
        row_cells[1].text = str(row['University Seat Number'])
        row_cells[2].text = col.replace('_Total', '')  # Remove _Total from subject name
        row_cells[3].text = str(int(max_score)) if pd.notna(max_score) else ""
        row_cells[4].text = str(int(row['Total_Full_Marks'])) if pd.notna(row['Total_Full_Marks']) else ""

# Save Word file
output_file = "Student_Results_Summary.docx"
doc.save(output_file)
print(f"Report saved as {output_file}")
'''

'''
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os


class ResultAnalyzer:
    def __init__(self, excel_file_path):
        self.excel_file = excel_file_path
        self.df = None
        self.df_clean = None
        self.subject_name_columns = []
        self.result_columns = []
        self.total_columns = []

    def load_and_prepare_data(self):
        """Load and clean data with comprehensive validation"""
        try:
            self.df = pd.read_excel(
                self.excel_file,
                sheet_name='Sheet1',
                dtype={
                    'University Seat Number': str,
                    'Student Name': str
                }
            )

            if not isinstance(self.df, pd.DataFrame):
                raise ValueError("Input is not a valid DataFrame")

            if self.df.empty:
                raise ValueError("Input Excel file is empty")

            self.df_clean = self.df.dropna(how='all')

            if self.df_clean.empty:
                raise ValueError("No valid data after cleaning")

            self._identify_columns()

            # Validate required columns exist
            required = ['Student Name', 'University Seat Number']
            missing = [col for col in required if col not in self.df_clean.columns]
            if missing:
                raise ValueError(f"Missing required columns: {', '.join(missing)}")

        except Exception as e:
            raise ValueError(f"Data loading failed: {str(e)}")

    def _identify_columns(self):
        """Identify columns with validation"""
        self.subject_name_columns = [
            col for col in self.df_clean.columns
            if isinstance(col, str) and col.endswith('_SubjectName')
        ]
        self.result_columns = [
            col for col in self.df_clean.columns
            if isinstance(col, str) and col.endswith('_Result')
        ]
        self.total_columns = [
            col for col in self.df_clean.columns
            if isinstance(col, str) and col.endswith('_Total')
        ]

        if not self.subject_name_columns:
            raise ValueError("No subject name columns found")

    @staticmethod
    def set_table_borders(table):
        """Set table borders in Word doc"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            borders.append(border)
        tblPr.append(borders)

    def generate_report(self, output_file=None):
        """Generate report with full error handling"""
        if not hasattr(self, 'df_clean') or self.df_clean is None:
            raise ValueError("Data not loaded. Call load_and_prepare_data() first")

        if self.df_clean.empty:
            raise ValueError("No data available to generate report")

        if not output_file:
            base_name = os.path.splitext(os.path.basename(self.excel_file))[0]
            output_file = f"{base_name}_Analysis_Report.docx"

        try:
            # Calculate Total Full Marks safely
            self.df_clean['Total_Full_Marks'] = self.df_clean[self.total_columns] \
                .apply(pd.to_numeric, errors='coerce') \
                .sum(axis=1, min_count=1)  # min_count requires at least 1 valid value

            doc = Document()
            self._add_title_section(doc)
            self._add_subject_summary(doc)
            self._add_results_by_subject(doc)
            self._add_top_students(doc)
            self._add_top_performers_by_subject(doc)

            doc.save(output_file)
            return output_file

        except Exception as e:
            raise ValueError(f"Report generation failed: {str(e)}")

    def _add_title_section(self, doc):
        """Add title section"""
        doc.add_heading("Student Results Summary Report", 0) \
            .alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_subject_summary(self, doc):
        """Add subject summary with validation"""
        doc.add_heading("Subject Code and Name Summary", level=1) \
            .alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = doc.add_table(rows=1, cols=4)
        self.set_table_borders(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "S.No"
        hdr_cells[1].text = "Subject Name"
        hdr_cells[2].text = "Subject Code"
        hdr_cells[3].text = "Total Appeared"

        serial_no = 1
        for sub_col in self.subject_name_columns:
            try:
                subject_code = sub_col.replace('_SubjectName', '')
                subject_name = self.df_clean[sub_col].dropna().astype(str).str.strip()
                subject_name = subject_name[subject_name != ''].unique()

                if len(subject_name) == 0:
                    continue

                subject_name = subject_name[0]
                result_col = f"{subject_code}_Result"

                appeared = 0
                if result_col in self.df_clean.columns:
                    appeared = self.df_clean[result_col].notna().sum()

                row_cells = table.add_row().cells
                row_cells[0].text = str(serial_no)
                row_cells[1].text = subject_name
                row_cells[2].text = subject_code
                row_cells[3].text = str(appeared)
                serial_no += 1

            except Exception:
                continue

    def _add_results_by_subject(self, doc):
        """Add results by subject with safe DataFrame handling"""
        doc.add_heading("Summary of Results by Subject", level=1) \
            .alignment = WD_ALIGN_PARAGRAPH.CENTER

        for col in self.result_columns:
            try:
                results = self.df_clean[col].dropna()
                if results.empty:
                    continue

                counts = results.value_counts()
                appeared_count = len(results)

                doc.add_heading(col, level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
                table = doc.add_table(rows=1, cols=3)
                self.set_table_borders(table)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Result'
                hdr_cells[1].text = 'Count'
                hdr_cells[2].text = 'Percentage'

                for result, count in counts.items():
                    percentage = (count / appeared_count) * 100
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(result)
                    row_cells[1].text = str(count)
                    row_cells[2].text = f"{percentage:.2f}%"

                row_cells = table.add_row().cells
                row_cells[0].text = 'Total Appeared'
                row_cells[1].text = str(appeared_count)
                row_cells[2].text = '—'

            except Exception:
                continue

    def _add_top_students(self, doc):
        """Add top students section with validation"""
        doc.add_heading("Top 10 Students by Total Full Marks", level=1) \
            .alignment = WD_ALIGN_PARAGRAPH.CENTER

        try:
            df_top = self.df_clean.dropna(subset=['Total_Full_Marks'])
            if df_top.empty:
                return

            top_10 = df_top.nlargest(10, 'Total_Full_Marks')[[
                'University Seat Number', 'Student Name', 'Total_Full_Marks'
            ]]

            if top_10.empty:
                return

            table = doc.add_table(rows=1, cols=3)
            self.set_table_borders(table)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'University Seat Number'
            hdr_cells[1].text = 'Student Name'
            hdr_cells[2].text = 'Total Full Marks'

            for _, row in top_10.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['University Seat Number'])
                row_cells[1].text = str(row['Student Name'])
                row_cells[2].text = str(int(row['Total_Full_Marks']))

        except Exception:
            return

    def _add_top_performers_by_subject(self, doc):
        """Completely fixed top performers section"""
        doc.add_heading("Top Performer in Each Subject", level=1) \
            .alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = doc.add_table(rows=1, cols=5)
        self.set_table_borders(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Student Name"
        hdr_cells[1].text = "University Seat Number"
        hdr_cells[2].text = "Subject Code"
        hdr_cells[3].text = "Subject Marks"
        hdr_cells[4].text = "Total Full Marks"

        for col in self.total_columns:
            try:
                subject_scores = pd.to_numeric(self.df_clean[col], errors='coerce')

                if not subject_scores.notna().any():
                    continue

                max_score = subject_scores.max()
                if pd.isna(max_score):
                    continue

                mask = subject_scores.notna() & (subject_scores == max_score)
                if not mask.any():
                    continue

                top_rows = self.df_clean.loc[mask]
                if top_rows.empty:
                    continue

                for _, row in top_rows.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row.get('Student Name', ''))
                    row_cells[1].text = str(row.get('University Seat Number', ''))
                    row_cells[2].text = col.replace('_Total', '')
                    row_cells[3].text = str(int(max_score)) if pd.notna(max_score) else ""
                    row_cells[4].text = str(int(row.get('Total_Full_Marks', 0))) \
                        if pd.notna(row.get('Total_Full_Marks')) else ""

            except Exception:
                continue


def analyze_results(excel_file_path, output_file=None):
    """Public interface with error handling"""
    try:
        analyzer = ResultAnalyzer(excel_file_path)
        analyzer.load_and_prepare_data()
        return analyzer.generate_report(output_file)
    except Exception as e:
        raise ValueError(f"Analysis failed: {str(e)}")
'''

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os


class ResultAnalyzer:
    def __init__(self, excel_file_path):
        self.excel_file = excel_file_path
        self.df = None
        self.df_clean = None
        self.subject_name_columns = []
        self.result_columns = []
        self.total_columns = []

    def load_and_prepare_data(self):
        """Load and clean data with comprehensive validation"""
        try:
            self.df = pd.read_excel(
                self.excel_file,
                sheet_name='Sheet1',
                dtype={
                    'University Seat Number': str,
                    'Student Name': str
                }
            )

            if not isinstance(self.df, pd.DataFrame):
                raise ValueError("Input is not a valid DataFrame")

            if self.df.empty:
                raise ValueError("Input Excel file is empty")

            self.df_clean = self.df.dropna(how='all')

            if self.df_clean.empty:
                raise ValueError("No valid data after cleaning")

            self._identify_columns()

            # Validate required columns exist
            required = ['Student Name', 'University Seat Number']
            missing = [col for col in required if col not in self.df_clean.columns]
            if missing:
                raise ValueError(f"Missing required columns: {', '.join(missing)}")

        except Exception as e:
            raise ValueError(f"Data loading failed: {str(e)}")

    def _identify_columns(self):
        """Identify columns with validation"""
        self.subject_name_columns = [
            col for col in self.df_clean.columns
            if isinstance(col, str) and col.endswith('_SubjectName')
        ]
        self.result_columns = [
            col for col in self.df_clean.columns
            if isinstance(col, str) and col.endswith('_Result')
        ]
        self.total_columns = [
            col for col in self.df_clean.columns
            if isinstance(col, str) and col.endswith('_Total')
        ]

        if not self.subject_name_columns:
            raise ValueError("No subject name columns found")

    @staticmethod
    def set_table_borders(table):
        """Set table borders in Word doc"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            borders.append(border)
        tblPr.append(borders)

    def _identify_failed_students(self):
        """Identify students who failed in any one subject"""
        failed_students = []
        total_students = len(self.df_clean)

        for _, student in self.df_clean.iterrows():
            student_failed = False
            failed_subjects = []

            for result_col in self.result_columns:
                result = student.get(result_col)
                # Check if result indicates failure (you might need to adjust these conditions based on your grading system)
                if pd.notna(result) and isinstance(result, str):
                    result_lower = result.strip().lower()
                    if result_lower in ['f', 'fail', 'failed', 'ab', 'absent']:
                        student_failed = True
                        subject_code = result_col.replace('_Result', '')
                        failed_subjects.append(subject_code)
                # Also check numeric failure conditions if applicable
                elif pd.notna(result) and isinstance(result, (int, float)):
                    if result < 35:  # Assuming 35 is passing marks
                        student_failed = True
                        subject_code = result_col.replace('_Result', '')
                        failed_subjects.append(subject_code)

            if student_failed:
                failed_students.append({
                    'Student Name': student.get('Student Name', ''),
                    'University Seat Number': student.get('University Seat Number', ''),
                    'Failed Subjects': ', '.join(failed_subjects),
                    'Total Failed Subjects': len(failed_subjects)
                })

        return failed_students, total_students

    def generate_report(self, output_file=None):
        """Generate report with full error handling"""
        if not hasattr(self, 'df_clean') or self.df_clean is None:
            raise ValueError("Data not loaded. Call load_and_prepare_data() first")

        if self.df_clean.empty:
            raise ValueError("No data available to generate report")

        if not output_file:
            base_name = os.path.splitext(os.path.basename(self.excel_file))[0]
            output_file = f"{base_name}_Analysis_Report.docx"

        try:
            # Calculate Total Full Marks safely
            self.df_clean['Total_Full_Marks'] = self.df_clean[self.total_columns] \
                .apply(pd.to_numeric, errors='coerce') \
                .sum(axis=1, min_count=1)  # min_count requires at least 1 valid value

            doc = Document()
            self._add_title_section(doc)
            self._add_subject_summary(doc)
            self._add_failed_students_summary(doc)  # NEW: Add failed students table
            self._add_results_by_subject(doc)
            self._add_top_students(doc)
            self._add_top_performers_by_subject(doc)

            doc.save(output_file)
            return output_file

        except Exception as e:
            raise ValueError(f"Report generation failed: {str(e)}")

    def _add_title_section(self, doc):
        """Add title section"""
        doc.add_heading("Student Results Summary Report", 0) \
            .alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_failed_students_summary(self, doc):
        """NEW: Add failed students summary table"""
        doc.add_heading("Failed Students Summary", level=1) \
            .alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Identify failed students
        failed_students, total_students = self._identify_failed_students()
        failed_count = len(failed_students)
        pass_count = total_students - failed_count

        # Create summary table for statistics
        summary_table = doc.add_table(rows=1, cols=4)
        self.set_table_borders(summary_table)
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Add header for summary table
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = "Total Students"
        hdr_cells[1].text = "Passed Students"
        hdr_cells[2].text = "Failed Students"
        hdr_cells[3].text = "Pass Percentage"

        # Add data to summary table
        if total_students > 0:
            pass_percentage = (pass_count / total_students) * 100
            fail_percentage = (failed_count / total_students) * 100
        else:
            pass_percentage = 0
            fail_percentage = 0

        data_cells = summary_table.add_row().cells
        data_cells[0].text = str(total_students)
        data_cells[1].text = f"{pass_count} ({pass_percentage:.2f}%)"
        data_cells[2].text = f"{failed_count} ({fail_percentage:.2f}%)"
        data_cells[3].text = f"{pass_percentage:.2f}%"

        doc.add_paragraph()  # Add some space

        # Add detailed failed students table
        if failed_students:
            doc.add_heading("Detailed Failed Students List", level=2)

            table = doc.add_table(rows=1, cols=4)
            self.set_table_borders(table)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "S.No"
            hdr_cells[1].text = "Student Name"
            hdr_cells[2].text = "University Seat Number"
            hdr_cells[3].text = "Failed Subjects (Count)"

            for i, student in enumerate(failed_students, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = student['Student Name']
                row_cells[2].text = student['University Seat Number']
                row_cells[3].text = f"{student['Failed Subjects']} ({student['Total Failed Subjects']})"
        else:
            doc.add_paragraph("No failed students found.", style='Intense Quote')

        doc.add_paragraph()  # Add some space between sections

    def _add_subject_summary(self, doc):
        """Add subject summary with validation"""
        doc.add_heading("Subject Code and Name Summary", level=1) \
            .alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = doc.add_table(rows=1, cols=4)
        self.set_table_borders(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "S.No"
        hdr_cells[1].text = "Subject Name"
        hdr_cells[2].text = "Subject Code"
        hdr_cells[3].text = "Total Appeared"

        serial_no = 1
        for sub_col in self.subject_name_columns:
            try:
                subject_code = sub_col.replace('_SubjectName', '')
                subject_name = self.df_clean[sub_col].dropna().astype(str).str.strip()
                subject_name = subject_name[subject_name != ''].unique()

                if len(subject_name) == 0:
                    continue

                subject_name = subject_name[0]
                result_col = f"{subject_code}_Result"

                appeared = 0
                if result_col in self.df_clean.columns:
                    appeared = self.df_clean[result_col].notna().sum()

                row_cells = table.add_row().cells
                row_cells[0].text = str(serial_no)
                row_cells[1].text = subject_name
                row_cells[2].text = subject_code
                row_cells[3].text = str(appeared)
                serial_no += 1

            except Exception:
                continue

    def _add_results_by_subject(self, doc):
        """Add results by subject with safe DataFrame handling"""
        doc.add_heading("Summary of Results by Subject", level=1) \
            .alignment = WD_ALIGN_PARAGRAPH.CENTER

        for col in self.result_columns:
            try:
                results = self.df_clean[col].dropna()
                if results.empty:
                    continue

                counts = results.value_counts()
                appeared_count = len(results)

                doc.add_heading(col, level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
                table = doc.add_table(rows=1, cols=3)
                self.set_table_borders(table)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Result'
                hdr_cells[1].text = 'Count'
                hdr_cells[2].text = 'Percentage'

                for result, count in counts.items():
                    percentage = (count / appeared_count) * 100
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(result)
                    row_cells[1].text = str(count)
                    row_cells[2].text = f"{percentage:.2f}%"

                row_cells = table.add_row().cells
                row_cells[0].text = 'Total Appeared'
                row_cells[1].text = str(appeared_count)
                row_cells[2].text = '—'

            except Exception:
                continue

    def _add_top_students(self, doc):
        """Add top students section with validation"""
        doc.add_heading("Top 10 Students by Total Full Marks\nNote: It Includes Backlog Papers As Well.", level=1) \
            .alignment = WD_ALIGN_PARAGRAPH.CENTER

        try:
            df_top = self.df_clean.dropna(subset=['Total_Full_Marks'])
            if df_top.empty:
                return

            top_10 = df_top.nlargest(10, 'Total_Full_Marks')[[
                'University Seat Number', 'Student Name', 'Total_Full_Marks'
            ]]

            if top_10.empty:
                return

            table = doc.add_table(rows=1, cols=3)
            self.set_table_borders(table)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'University Seat Number'
            hdr_cells[1].text = 'Student Name'
            hdr_cells[2].text = 'Total Full Marks'

            for _, row in top_10.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['University Seat Number'])
                row_cells[1].text = str(row['Student Name'])
                row_cells[2].text = str(int(row['Total_Full_Marks']))

        except Exception:
            return

    def _add_top_performers_by_subject(self, doc):
        """Completely fixed top performers section"""
        doc.add_heading("Top Performer in Each Subject", level=1) \
            .alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = doc.add_table(rows=1, cols=5)
        self.set_table_borders(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Student Name"
        hdr_cells[1].text = "University Seat Number"
        hdr_cells[2].text = "Subject Code"
        hdr_cells[3].text = "Subject Marks"
        hdr_cells[4].text = "Total Full Marks"

        for col in self.total_columns:
            try:
                subject_scores = pd.to_numeric(self.df_clean[col], errors='coerce')

                if not subject_scores.notna().any():
                    continue

                max_score = subject_scores.max()
                if pd.isna(max_score):
                    continue

                mask = subject_scores.notna() & (subject_scores == max_score)
                if not mask.any():
                    continue

                top_rows = self.df_clean.loc[mask]
                if top_rows.empty:
                    continue

                for _, row in top_rows.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row.get('Student Name', ''))
                    row_cells[1].text = str(row.get('University Seat Number', ''))
                    row_cells[2].text = col.replace('_Total', '')
                    row_cells[3].text = str(int(max_score)) if pd.notna(max_score) else ""
                    row_cells[4].text = str(int(row.get('Total_Full_Marks', 0))) \
                        if pd.notna(row.get('Total_Full_Marks')) else ""

            except Exception:
                continue


def analyze_results(excel_file_path, output_file=None):
    """Public interface with error handling"""
    try:
        analyzer = ResultAnalyzer(excel_file_path)
        analyzer.load_and_prepare_data()
        return analyzer.generate_report(output_file)
    except Exception as e:
        raise ValueError(f"Analysis failed: {str(e)}")



def test_with_existing_file():
    """Test with your existing Excel file"""
    try:
        # Replace with your actual file path
        excel_file_path = "results.xlsx"

        # Generate report
        output_file = analyze_results(excel_file_path)

        print(f"✓ Report generated: {output_file}")
        print("Open the Word document to check the failed students table")

    except Exception as e:
        print(f"✗ Error: {e}")


# Run the test
test_with_existing_file()




